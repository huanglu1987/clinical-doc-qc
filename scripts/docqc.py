#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docqc.py —— 临床试验申报资料「格式层 + 结构层」核查工具

■ 为什么必须有这个工具
  AI 读到的文档内容是纯文本，**不携带字体、字号、行距、页边距信息**。
  仅凭"读一遍"判断排版是否合规，只会得到编造的结论。本工具直接解析
  .docx 的 OOXML 与 .pdf 的对象结构，把格式属性作为**可核查的事实**取出来，
  再交给规则集判定——事实归工具，判断归规则。

■ 三条能力线
  1. 格式合规   ：字体/字号/行距/颜色/页边距/纸张，对照 CDE 或 FDA 规则集
  2. 单文件一致性：表图编号是否连续、交叉引用是否指向存在的对象、
                   缩写定义是否前后冲突、同级样式是否统一、标题是否跳级
  3. 跨文件一致性：同一套资料（方案 / ICF / IB）之间格式是否统一

■ 用法
  python3 docqc.py check   <文件...> --profile cde|fda [--json]
  python3 docqc.py extract <文件...> [--json]          # 只出事实，不判定
  python3 docqc.py compare <文件A> <文件B> [...]        # 跨文件一致性

■ 依赖
  python-docx / pypdf / PyMuPDF(fitz)

■ 边界（务必知悉）
  本工具只覆盖「格式层 + 结构层」。
  内容前后矛盾、逻辑是否自洽、语言是否自然 属「文本层」，
  必须按 SKILL.md 的文本层流程由 AI 另行核查，工具不负责也判不了。
"""

import argparse
import json
import os
import re
import sys
from collections import Counter, defaultdict

EMU_PER_CM = 360000.0
EMU_PER_INCH = 914400.0
PT_PER_INCH = 72.0

# ── 中文字号 ↔ 磅值（输出时把 12.0pt 还原成「小四」，便于对照规范原文）──
CN_SIZE_NAME = {
    42.0: "初号", 36.0: "小初", 26.0: "一号", 24.0: "小一", 22.0: "二号",
    18.0: "小二", 16.0: "三号", 15.0: "小三", 14.0: "四号", 12.0: "小四",
    10.5: "五号", 9.0: "小五", 7.5: "六号", 6.5: "小六",
}


def _pt_label(pt):
    if pt is None:
        return "—"
    name = CN_SIZE_NAME.get(round(float(pt), 1))
    return f"{pt}pt（{name}）" if name else f"{pt}pt"


# ══════════════════════════════════════════════════════════════
#  规则集
# ══════════════════════════════════════════════════════════════
# CDE：《药品注册申报资料格式体例与整理规范》（CDE，2020-07 发布，2020-10-01 施行）
#      具体数值按本部门确认口径固化，详见 references/cde-format-spec.md
CDE_SPEC = {
    "label": "CDE（中国）",
    "cn_font": "宋体",
    "latin_font": "Times New Roman",
    "body_size_pt": 12.0,        # 小四
    "table_min_size_pt": 10.5,   # 五号
    "line_spacing": 1.5,         # 1.5 倍行距
    "font_color_hex": "000000",  # 黑色
    "paper": {"name": "A4", "width_cm": 21.0, "height_cm": 29.7, "tol_cm": 0.3},
    "margins_cm": {
        "portrait":  {"left": 2.5, "top": 2.0, "right": 1.0, "bottom": 1.0},
        "landscape": {"top": 2.5, "right": 2.0, "left": 1.0, "bottom": 1.0},
    },
}

# FDA：PDF Specifications for FDA Regulatory Submissions
#      详见 references/fda-format-spec.md
FDA_SPEC = {
    "label": "FDA（美国）",
    "latin_fonts_ok": ["Times New Roman", "Arial"],
    "body_size_pt": 12.0,
    "table_min_size_pt": 10.0,
    "table_hard_min_pt": 9.0,
    "margin_in": {"all": 1.0, "hard_left": 0.75, "hard_other": 0.375},
    "pdf_version_range": (1.4, 1.7),
    "bookmark_max_depth": 4,
    "bookmark_required_pages": 5,
}


def _norm_font(name):
    return (name or "").replace(" ", "").replace("-", "").lower()


def _font_matches(actual, expected):
    """字体名归一化比对，容忍 TimesNewRomanPSMT / Arial-BoldMT 这类 PDF 内部名。"""
    a, e = _norm_font(actual), _norm_font(expected)
    return bool(a) and (a == e or a.startswith(e) or e in a)


# ══════════════════════════════════════════════════════════════
#  通用
# ══════════════════════════════════════════════════════════════

def _fail(msg):
    print(f"[错误] {msg}", file=sys.stderr)
    sys.exit(1)


def _file_size_mb(path):
    return round(os.path.getsize(path) / (1024 * 1024), 2)


def _has_cjk(text):
    return any("一" <= ch <= "鿿" for ch in text)


# ══════════════════════════════════════════════════════════════
#  DOCX 事实提取
# ══════════════════════════════════════════════════════════════

def _style_chain(style):
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        yield style
        try:
            style = style.base_style
        except Exception:
            return


def _rfonts(el, attr):
    from docx.oxml.ns import qn
    try:
        rPr = el.rPr
        if rPr is None:
            return None
        rFonts = rPr.rFonts
        if rFonts is None:
            return None
        return rFonts.get(qn("w:" + attr))
    except Exception:
        return None


def _doc_defaults(doc):
    """styles.xml 的 docDefaults：所有未显式设置字体的文字最终继承它。"""
    from docx.oxml.ns import qn
    out = {"latin": None, "eastasia": None, "size_pt": None}
    try:
        dd = doc.styles.element.find(qn("w:docDefaults"))
        rpd = dd.find(qn("w:rPrDefault")) if dd is not None else None
        rPr = rpd.find(qn("w:rPr")) if rpd is not None else None
        if rPr is None:
            return out
        rf = rPr.find(qn("w:rFonts"))
        if rf is not None:
            out["latin"] = rf.get(qn("w:ascii"))
            out["eastasia"] = rf.get(qn("w:eastAsia"))
        sz = rPr.find(qn("w:sz"))
        if sz is not None and sz.get(qn("w:val")):
            out["size_pt"] = float(sz.get(qn("w:val"))) / 2.0   # w:sz 单位是半磅
    except Exception:
        pass
    return out


def _resolve_run(run, para, defaults):
    """解析一个 run 最终生效的中/西文字体、字号、颜色（run → 样式链 → 文档默认）。"""
    latin = run.font.name or _rfonts(run._element, "ascii")
    east = _rfonts(run._element, "eastAsia")
    size = run.font.size.pt if run.font.size is not None else None

    color = None
    try:
        if run.font.color is not None and run.font.color.rgb is not None:
            color = str(run.font.color.rgb)
    except Exception:
        pass

    if latin is None or east is None or size is None:
        try:
            for st in _style_chain(para.style):
                latin = latin or st.font.name or _rfonts(st.element, "ascii")
                east = east or _rfonts(st.element, "eastAsia")
                if size is None and st.font.size is not None:
                    size = st.font.size.pt
                if latin and east and size:
                    break
        except Exception:
            pass

    return {
        "latin": latin or defaults["latin"],
        "eastasia": east or defaults["eastasia"],
        "size_pt": round(size, 1) if size is not None else defaults["size_pt"],
        "color": color,
    }


def _resolve_line_spacing(para):
    """行距：段落 → 样式链。返回 (倍数, 是否为固定值pt)。"""
    try:
        ls = para.paragraph_format.line_spacing
        if ls is not None:
            return (round(float(ls), 2), False) if isinstance(ls, float) else (round(ls.pt, 1), True)
    except Exception:
        pass
    try:
        for st in _style_chain(para.style):
            ls = st.paragraph_format.line_spacing
            if ls is not None:
                return (round(float(ls), 2), False) if isinstance(ls, float) else (round(ls.pt, 1), True)
    except Exception:
        pass
    return (None, False)


# ── 结构层：编号 / 交叉引用 / 缩写 ────────────────────────────
RE_CAPTION = re.compile(r"^\s*(表|图|附录|附件|Table|Figure|Appendix)\s*([0-9]+(?:[-.—][0-9]+)*|[A-Z])(?![0-9])")
# 正文中出现的任意「表N / 图N / Table N」都按交叉引用处理——
# 不能只认「见」字后面那一个，"详见表 1 与表 5" 里的表 5 同样是引用。
RE_ANY_REF = re.compile(r"(表|图|附录|附件|Table|Figure|Appendix)\s*([0-9]+(?:[-.—][0-9]+)*)", re.I)
RE_SEC_XREF = re.compile(r"(?:见|参见|详见)\s*第?\s*([0-9]+(?:\.[0-9]+)*)\s*(?:节|章|部分)?")
RE_ABBR_PAREN = re.compile(r"[（(]\s*([A-Z][A-Za-z0-9]{1,7})\s*[)）]")
RE_ABBR_USE = re.compile(r"(?<![A-Za-z])([A-Z]{2,6})(?![A-Za-z])")

ABBR_STOP = {"CDE", "FDA", "NMPA", "EMA", "ICH", "GCP", "PDF", "CRF", "CRO", "SMO",
             "OK", "ID", "NA", "II", "III", "IV", "VI", "PI", "US", "EU", "CN", "AND", "THE", "FOR"}


def _abbr_defs_in(text):
    """抽取「全称（缩写）」定义对。

    不能直接用一个贪婪正则往前吃——"本研究评价不良事件（AE）" 会把整句当成全称。
    正确做法是从括号往前走，遇标点即停，中文取末尾连续汉字、英文取末尾若干单词。
    """
    out = []
    for m in RE_ABBR_PAREN.finditer(text):
        ab = m.group(1)
        if ab in ABBR_STOP:
            continue
        seg = re.split(r"[，。；：、,;:!?！？\n\t（(]", text[:m.start()])[-1].strip()
        if not seg:
            continue
        cjk = re.search(r"([一-鿿]{2,12})$", seg)
        if cjk:
            full = cjk.group(1)
        else:
            words = re.findall(r"[A-Za-z][A-Za-z\-]*", seg)
            full = " ".join(words[-6:]) if words else ""
        if len(full) >= 2:
            out.append((full, ab))
    return out


# ── 目录（TOC）提取 ──────────────────────────────────────────
RE_TOC_STYLE = re.compile(r"^(toc|目录)\s*\d*$", re.I)
# "1 研究背景........5" / "1 研究背景\t5" / "1 研究背景   5"
RE_TOC_LINE = re.compile(r"^(.+?)[\s.．·……\t]{2,}(\d{1,4})\s*$")


def _norm_title(s):
    return re.sub(r"[\s.．·……]+", "", (s or "")).strip()


def _extract_toc_docx(doc):
    """从 .docx 抽取目录条目。

    ⚠ 关键：.docx 里目录的页码是**域的缓存值**，Word 打开或转 PDF 时会重算。
    未渲染状态下读到的数字可能早已过期，**据此判定"页码错误"是不可靠的**。
    因此这里只把页码原样带出来标为 stale（供人工参考），判定交给标题比对。
    """
    entries = []
    for p in doc.paragraphs:
        sname = p.style.name if p.style is not None else ""
        if not RE_TOC_STYLE.match(sname or ""):
            continue
        raw = p.text.strip()
        if not raw:
            continue
        title, page = raw, None
        if "\t" in raw:
            parts = [x for x in raw.split("\t") if x.strip()]
            if len(parts) >= 2 and parts[-1].strip().isdigit():
                title, page = "\t".join(parts[:-1]).strip(), int(parts[-1].strip())
        if page is None:
            m = RE_TOC_LINE.match(raw)
            if m:
                title, page = m.group(1).strip(), int(m.group(2))
        lvl = "".join(ch for ch in sname if ch.isdigit())
        entries.append({"title": title, "cached_page": page,
                        "level": int(lvl) if lvl.isdigit() else None})
    return entries


def _extract_toc_pdf(fdoc, max_entries=80):
    """从 PDF 抽取「印在纸面上的目录」，并定位每个标题实际出现在第几页。

    PDF 已渲染，页码是实的，因此这里可以真正比对。
    注意正文页码与 PDF 物理页序常有固定偏移（封面/扉页），
    所以采用「众数偏移」判定：只报偏离众数的条目，不把整体偏移当错误。
    """
    toc_pages, entries = [], []
    scan = min(len(fdoc), 12)
    for pno in range(scan):
        lines = [l.strip() for l in (fdoc[pno].get_text("text") or "").splitlines() if l.strip()]
        hits = [l for l in lines if RE_TOC_LINE.match(l)]
        if len(hits) >= 4:                     # 一页里有 4 行以上「标题…页码」才算目录页
            toc_pages.append(pno)
            for l in hits:
                m = RE_TOC_LINE.match(l)
                entries.append({"title": m.group(1).strip(), "stated_page": int(m.group(2))})
    if not entries:
        return {"found": False, "entries": [], "toc_physical_pages": []}

    entries = entries[:max_entries]
    body_start = max(toc_pages) + 1 if toc_pages else 0
    for e in entries:
        needle = e["title"]
        if len(needle) < 3:
            e["actual_physical_page"] = None
            continue
        found = None
        for pno in range(body_start, len(fdoc)):
            try:
                if fdoc[pno].search_for(needle[:60], quads=False):
                    found = pno + 1            # 1-based 物理页
                    break
            except Exception:
                break
        e["actual_physical_page"] = found

    offs = [e["actual_physical_page"] - e["stated_page"]
            for e in entries if e.get("actual_physical_page")]
    modal = Counter(offs).most_common(1)[0][0] if offs else None
    for e in entries:
        ap = e.get("actual_physical_page")
        e["offset"] = (ap - e["stated_page"]) if ap else None
        e["deviates"] = bool(modal is not None and e["offset"] is not None
                             and e["offset"] != modal)
    return {"found": True, "entries": entries,
            "toc_physical_pages": [p + 1 for p in toc_pages],
            "modal_offset": modal,
            "resolved": sum(1 for e in entries if e.get("actual_physical_page"))}


# ── 编辑痕迹与排版细节（定稿前校对高发项）──────────────────
RE_FIELD_ERR = re.compile(
    r"错误[!！]\s*未找到引用源|错误[!！]\s*超链接引用无效|"
    r"Error!\s*Reference source not found|Error!\s*Bookmark not defined", re.I)
RE_HALF_PUNCT = re.compile(r"[一-鿿]\s*([,;!?])")          # 中文字后紧跟半角标点
RE_DATE_STYLES = {
    "YYYY-MM-DD": re.compile(r"\d{4}-\d{1,2}-\d{1,2}"),
    "YYYY年M月D日": re.compile(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日"),
    "YYYY.MM.DD": re.compile(r"\d{4}\.\d{1,2}\.\d{1,2}"),
    "YYYY/MM/DD": re.compile(r"\d{4}/\d{1,2}/\d{1,2}"),
    "DD/MM/YYYY 或 MM/DD/YYYY": re.compile(r"(?<!\d)\d{1,2}/\d{1,2}/\d{2,4}"),
}


def _extract_editorial(doc, path):
    """抽取定稿前最该查、但肉眼最容易漏的几类问题。"""
    from docx.oxml.ns import qn
    out = {}

    # ① 修订痕迹（发出去带修订痕迹属于事故）
    body = doc.element.body
    cnt = lambda tag: len(body.findall(".//" + qn(tag)))
    out["tracked_changes"] = {
        "insertions": cnt("w:ins"), "deletions": cnt("w:del"),
        "moves": cnt("w:moveFrom") + cnt("w:moveTo"),
        "format_changes": cnt("w:rPrChange") + cnt("w:pPrChange"),
    }

    # ② 批注残留
    ncom = 0
    try:
        import zipfile
        with zipfile.ZipFile(path) as z:
            if "word/comments.xml" in z.namelist():
                ncom = z.read("word/comments.xml").decode("utf-8", "ignore").count("<w:comment ")
    except Exception:
        pass
    out["comments"] = ncom

    # ③ 页眉页脚（版本号/方案编号改了但页眉没改，是高发问题）
    heads, foots = [], []
    for s in doc.sections:
        try:
            heads.append(" / ".join(p.text.strip() for p in s.header.paragraphs if p.text.strip()))
            foots.append(" / ".join(p.text.strip() for p in s.footer.paragraphs if p.text.strip()))
        except Exception:
            pass
    out["headers"] = heads
    out["footers"] = foots

    # ④ 空段落与空表格单元格（结构性编辑后的残留）
    empties, run = 0, 0
    for p in doc.paragraphs:
        if p.text.strip():
            run = 0
        else:
            run += 1
            if run >= 2:
                empties += 1
    empty_cells = 0
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if not c.text.strip():
                    empty_cells += 1
    out["consecutive_empty_paragraphs"] = empties
    out["empty_table_cells"] = empty_cells

    # ⑤ 文档属性（可能带出上一版作者/单位信息）
    try:
        cp = doc.core_properties
        out["properties"] = {"author": cp.author or None,
                             "last_modified_by": cp.last_modified_by or None,
                             "title": cp.title or None,
                             "category": cp.category or None}
    except Exception:
        out["properties"] = {}
    return out


def _scan_text_issues(paras_text, is_cjk_doc):
    """纯文本层面的排版细节：域错误、日期格式混用、半角标点、括号配对。"""
    out = {"field_errors": [], "date_styles": {}, "half_punct": [], "unbalanced": []}
    for t in paras_text:
        if not t:
            continue
        if RE_FIELD_ERR.search(t):
            out["field_errors"].append(t[:60])
        if is_cjk_doc:
            # 同一段里多处命中只记一条，避免上下文窗口重叠、输出刷屏
            marks = RE_HALF_PUNCT.findall(t)
            if marks:
                out["half_punct"].append(
                    {"text": t[:60], "marks": "".join(sorted(set(marks))), "count": len(marks)})
        for pair, (l, r) in {"（）": ("（", "）"), "()": ("(", ")"),
                             "【】": ("【", "】")}.items():
            if t.count(l) != t.count(r):
                out["unbalanced"].append({"pair": pair, "text": t[:60]})
    joined = "\n".join(paras_text)
    for name, rx in RE_DATE_STYLES.items():
        n = len(rx.findall(joined))
        if n:
            out["date_styles"][name] = n
    out["field_errors"] = out["field_errors"][:10]
    out["half_punct"] = out["half_punct"][:10]
    out["unbalanced"] = out["unbalanced"][:10]
    return out


def _dedupe_fullnames(fulls):
    """归并同一缩写的多个候选全称。

    中文里全称的左边界没有语法标记（"本研究评价不良事件（AE）"），抽取难免多带几个字。
    若一个候选是另一个的后缀（不良事件 ⊂ 本研究评价不良事件），视为同一定义，取最短的那个，
    避免因抽取误差报出假冲突；真正不同的定义（不良事件 vs 不良反应）仍会保留并报警。
    """
    out = []
    for x in sorted(set(fulls), key=len):
        if not any(x.endswith(y) or y.endswith(x) for y in out):
            out.append(x)
    return out


def _analyze_structure(paras_text):
    """从全文段落文本里抽取编号、交叉引用、缩写定义，供一致性判定。"""
    captions = defaultdict(list)   # {"表": ["1","2",...]}
    xrefs = defaultdict(set)
    sec_refs = set()
    abbr_defs = defaultdict(set)   # {"AE": {"不良事件", ...}}
    abbr_uses = Counter()
    sec_numbers = set()

    for t in paras_text:
        if not t:
            continue
        rest = t
        m = RE_CAPTION.match(t)
        if m:
            captions[m.group(1)].append(m.group(2))
            rest = t[m.end():]        # 标题自身不算引用，只扫描其余部分
        for kind, num in RE_ANY_REF.findall(rest):
            k = {"table": "Table", "figure": "Figure", "appendix": "Appendix"}.get(kind.lower(), kind)
            xrefs[k].add(num)
        for sec in RE_SEC_XREF.findall(t):
            sec_refs.add(sec)
        for full, ab in _abbr_defs_in(t):
            abbr_defs[ab].add(full)
        for ab in RE_ABBR_USE.findall(t):
            if ab not in ABBR_STOP:
                abbr_uses[ab] += 1
        ms = re.match(r"^\s*([0-9]+(?:\.[0-9]+)*)\s+\S", t)
        if ms:
            sec_numbers.add(ms.group(1))

    return {
        "captions": {k: v for k, v in captions.items()},
        "xrefs": {k: sorted(v) for k, v in xrefs.items()},
        "section_refs": sorted(sec_refs),
        "section_numbers": sorted(sec_numbers),
        "abbr_defs": {k: sorted(v) for k, v in abbr_defs.items()},
        "abbr_uses": abbr_uses.most_common(),
    }


def extract_docx(path):
    try:
        from docx import Document
    except ImportError:
        _fail("缺少 python-docx，请执行：pip3 install python-docx")

    doc = Document(path)
    defaults = _doc_defaults(doc)

    sections = []
    for s in doc.sections:
        def cm(v):
            return round(v / EMU_PER_CM, 2) if v is not None else None
        w, h = cm(s.page_width), cm(s.page_height)
        sections.append({
            "orientation": "landscape" if (w and h and w > h) else "portrait",
            "page_width_cm": w, "page_height_cm": h,
            "left_cm": cm(s.left_margin), "right_cm": cm(s.right_margin),
            "top_cm": cm(s.top_margin), "bottom_cm": cm(s.bottom_margin),
        })

    latin_c, east_c, size_c, color_c = Counter(), Counter(), Counter(), Counter()
    ls_c = Counter()
    style_fonts = defaultdict(lambda: {"latin": set(), "eastasia": set(), "sizes": set()})
    headings, paras_text = [], []
    body_chars = 0
    cjk = False

    for para in doc.paragraphs:
        text = para.text.strip()
        paras_text.append(text)
        if text:
            if _has_cjk(text):
                cjk = True
            sp, is_exact = _resolve_line_spacing(para)
            if sp is not None:
                ls_c[("固定值" if is_exact else "倍数", sp)] += 1

        sname = para.style.name if para.style is not None else "(none)"
        for run in para.runs:
            if not run.text.strip():
                continue
            f = _resolve_run(run, para, defaults)
            n = len(run.text)
            body_chars += n
            if f["latin"]:
                latin_c[f["latin"]] += n
                style_fonts[sname]["latin"].add(f["latin"])
            if f["eastasia"]:
                east_c[f["eastasia"]] += n
                style_fonts[sname]["eastasia"].add(f["eastasia"])
            if f["size_pt"]:
                size_c[f["size_pt"]] += n
                style_fonts[sname]["sizes"].add(f["size_pt"])
            color_c[f["color"] or "(继承/自动)"] += n

        if sname.lower().startswith("heading") and text:
            lvl = "".join(ch for ch in sname if ch.isdigit())
            fr = para.runs[0] if para.runs else None
            hf = _resolve_run(fr, para, defaults) if fr is not None else {}
            headings.append({
                "level": int(lvl) if lvl.isdigit() else None,
                "style": sname, "text": text[:80],
                "latin": hf.get("latin"), "eastasia": hf.get("eastasia"),
                "size_pt": hf.get("size_pt"),
            })

    # 表格：字号单独统计（CDE/FDA 对表格字号都有单独下限）
    tbl_sizes = Counter()
    tbl_count = 0
    for tbl in doc.tables:
        tbl_count += 1
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        paras_text.append(para.text.strip())
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        f = _resolve_run(run, para, defaults)
                        if f["size_pt"]:
                            tbl_sizes[f["size_pt"]] += len(run.text)

    skips, prev = [], None
    for h in headings:
        lv = h["level"]
        if lv is None:
            continue
        if prev is not None and lv > prev + 1:
            skips.append({"from": prev, "to": lv, "text": h["text"]})
        prev = lv

    return {
        "file": os.path.basename(path), "path": os.path.abspath(path),
        "type": "docx", "file_size_mb": _file_size_mb(path),
        "language": "zh" if cjk else "en",
        "doc_defaults": defaults,
        "sections": sections,
        "fonts": {
            "latin_by_chars": latin_c.most_common(),
            "eastasia_by_chars": east_c.most_common(),
            "size_pt_by_chars": sorted(size_c.items(), key=lambda x: -x[1]),
            "color_by_chars": color_c.most_common(),
            "dominant_latin": latin_c.most_common(1)[0][0] if latin_c else None,
            "dominant_eastasia": east_c.most_common(1)[0][0] if east_c else None,
            "dominant_size_pt": size_c.most_common(1)[0][0] if size_c else None,
        },
        "line_spacing": [{"kind": k[0], "value": k[1], "paragraphs": v}
                         for k, v in ls_c.most_common()],
        "styles": {k: {"latin": sorted(v["latin"]), "eastasia": sorted(v["eastasia"]),
                       "sizes_pt": sorted(v["sizes"])}
                   for k, v in sorted(style_fonts.items())},
        "headings": {"count": len(headings), "level_skips": skips, "items": headings[:200]},
        "tables": {"count": tbl_count,
                   "size_pt_by_chars": sorted(tbl_sizes.items(), key=lambda x: -x[1]),
                   "min_size_pt": min(tbl_sizes) if tbl_sizes else None},
        "structure": _analyze_structure(paras_text),
        "toc": {"entries": _extract_toc_docx(doc), "page_numbers_verifiable": False},
        "editorial": dict(_extract_editorial(doc, path),
                          **_scan_text_issues(paras_text, cjk)),
        "body_chars": body_chars,
    }


# ══════════════════════════════════════════════════════════════
#  PDF 事实提取
# ══════════════════════════════════════════════════════════════

def extract_pdf(path):
    out = {"file": os.path.basename(path), "path": os.path.abspath(path),
           "type": "pdf", "file_size_mb": _file_size_mb(path)}
    try:
        from pypdf import PdfReader
    except ImportError:
        _fail("缺少 pypdf，请执行：pip3 install pypdf")

    reader = PdfReader(path)
    out["pdf_version"] = (getattr(reader, "pdf_header", "") or "").replace("%PDF-", "").strip() or None
    out["encrypted"] = bool(reader.is_encrypted)
    out["page_count"] = len(reader.pages)

    def depth(items, d=1):
        cnt = mx = 0
        for it in items:
            if isinstance(it, list):
                c2, d2 = depth(it, d + 1)
                cnt += c2
                mx = max(mx, d2)
            else:
                cnt += 1
                mx = max(mx, d)
        return cnt, mx

    try:
        bc, bd = depth(reader.outline)
    except Exception:
        bc, bd = 0, 0
    out["bookmarks"] = {"count": bc, "max_depth": bd}

    has_js = False
    try:
        root = reader.trailer["/Root"]
        if "/Names" in root and "/JavaScript" in root["/Names"]:
            has_js = True
        oa = root.get("/OpenAction")
        if isinstance(oa, dict) and oa.get("/S") == "/JavaScript":
            has_js = True
    except Exception:
        pass
    out["has_javascript"] = has_js

    try:
        import fitz
    except ImportError:
        out["_note"] = "未安装 PyMuPDF，字体与页边距未提取（pip3 install pymupdf）"
        return out

    d = fitz.open(path)
    try:
        out["toc"] = dict(_extract_toc_pdf(d), page_numbers_verifiable=True)
    except Exception:
        out["toc"] = {"found": False, "entries": [], "page_numbers_verifiable": True}
    fonts, sizes, dims = Counter(), Counter(), Counter()
    emb = {"embedded": 0, "not_embedded": 0}
    margins, text_chars = [], 0
    paras_text = []

    for pno in range(len(d)):
        page = d[pno]
        r = page.rect
        dims[(round(r.width / PT_PER_INCH, 2), round(r.height / PT_PER_INCH, 2))] += 1
        for f in page.get_fonts(full=True):
            fonts[(f[3] or "").split("+")[-1]] += 1
            emb["not_embedded" if f[1] in ("n/a", "", None) else "embedded"] += 1
        txt = page.get_text("text") or ""
        text_chars += len(txt.strip())
        paras_text.extend(l.strip() for l in txt.splitlines() if l.strip())
        try:
            dd = page.get_text("dict")
            x0 = y0 = 1e9
            x1 = y1 = -1e9
            for blk in dd.get("blocks", []):
                if blk.get("type") != 0:
                    continue
                bx0, by0, bx1, by1 = blk["bbox"]
                x0, y0, x1, y1 = min(x0, bx0), min(y0, by0), max(x1, bx1), max(y1, by1)
                for line in blk.get("lines", []):
                    for sp in line.get("spans", []):
                        s = round(sp.get("size", 0), 1)
                        if s:
                            sizes[s] += len(sp.get("text", ""))
            if x1 > x0:
                margins.append({
                    "left_in": round(x0 / PT_PER_INCH, 2),
                    "right_in": round((r.width - x1) / PT_PER_INCH, 2),
                    "top_in": round(y0 / PT_PER_INCH, 2),
                    "bottom_in": round((r.height - y1) / PT_PER_INCH, 2)})
        except Exception:
            pass
    d.close()

    mins = {k: min(m[k] for m in margins) for k in
            ("left_in", "right_in", "top_in", "bottom_in")} if margins else None

    out["fonts"] = {"used": fonts.most_common(), "embedding": emb,
                    "size_pt_by_chars": sorted(sizes.items(), key=lambda x: -x[1])[:20],
                    "dominant_size_pt": sizes.most_common(1)[0][0] if sizes else None}
    out["page_sizes_in"] = [{"size": list(k), "pages": v} for k, v in dims.most_common()]
    out["min_text_margins_in"] = mins
    out["text_searchable"] = text_chars > 50
    out["text_chars"] = text_chars
    out["language"] = "zh" if any(_has_cjk(t) for t in paras_text[:400]) else "en"
    out["structure"] = _analyze_structure(paras_text)
    out["editorial"] = _scan_text_issues(paras_text, out["language"] == "zh")
    return out


def extract(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".doc":
        _fail(f"{path}：不支持旧版 .doc，请先在 Word 里另存为 .docx")
    _fail(f"{path}：仅支持 .docx / .pdf")


# ══════════════════════════════════════════════════════════════
#  结构层一致性（单文件内部，CDE / FDA 通用）
# ══════════════════════════════════════════════════════════════

def _check_toc(d):
    """目录与正文一致性。

    ⚠ .docx 与 .pdf 能判的东西不同，必须分开处理：
      - .docx：目录页码是域的缓存值，未渲染时读到的可能早已过期。
               据此报"页码错误"会得出与事实相反的结论（本项目方法论第 8 条的真实反例）。
               因此 docx 只比对标题，页码一律不判，改为提示转 PDF 后复核。
      - .pdf ：页码是实的，可以真正比对；但正文页码与物理页序常有固定偏移
               （封面/扉页），故按"众数偏移"判定，只报偏离众数的条目。
    """
    f = []
    add = lambda lv, item, msg, fact=None: f.append(
        {"level": lv, "item": item, "message": msg, "fact": fact})
    toc = d.get("toc") or {}
    entries = toc.get("entries") or []
    if not entries:
        return f

    # ── 标题比对（两种格式都能做）──
    if d["type"] == "docx":
        heads = {_norm_title(h["text"]) for h in (d.get("headings", {}).get("items") or [])}
        if heads:
            missing = [e["title"] for e in entries
                       if _norm_title(e["title"]) and _norm_title(e["title"]) not in heads]
            if missing:
                add("ERROR", "目录与正文标题不符",
                    f"目录中有 {len(missing)} 条在正文标题里找不到对应项："
                    f"{'；'.join(missing[:5])}{' 等' if len(missing) > 5 else ''}",
                    missing[:5])
            toc_set = {_norm_title(e["title"]) for e in entries}
            absent = [h["text"] for h in (d.get("headings", {}).get("items") or [])
                      if h.get("level") in (1, 2) and _norm_title(h["text"]) not in toc_set]
            if absent:
                add("WARN", "正文标题未进目录",
                    f"{len(absent)} 个一/二级标题未出现在目录中："
                    f"{'；'.join(absent[:5])}{' 等' if len(absent) > 5 else ''}；"
                    f"通常是目录未更新（在 Word 里右键目录→更新域→更新整个目录）",
                    absent[:5])

        cached = [e for e in entries if e.get("cached_page") is not None]
        if cached:
            add("MANUAL", "目录页码（无法自动判定）",
                f"已读到 {len(cached)} 条目录页码，但 .docx 里页码是域的缓存值、"
                f"Word 打开或转 PDF 时会重算，未渲染状态下判定页码对错不可靠。"
                f"请在 Word 中「更新整个目录」后导出 PDF，再对 PDF 跑一次本工具核对页码。",
                len(cached))

    # ── 页码比对（仅 PDF）──
    else:
        dev = [e for e in entries if e.get("deviates")]
        unresolved = [e for e in entries if not e.get("actual_physical_page")]
        if dev:
            desc = "；".join(
                f"「{e['title'][:24]}」目录标 {e['stated_page']} 页、"
                f"实际在第 {e['actual_physical_page']} 页" for e in dev[:5])
            add("ERROR", "目录页码与正文不符",
                f"{len(dev)} 条目录页码偏离整体偏移（整体偏移 {toc.get('modal_offset')} 页）：{desc}",
                [e["title"] for e in dev[:5]])
        if unresolved:
            add("WARN", "目录条目未能定位",
                f"{len(unresolved)} 条目录标题在正文中未检索到，可能是标题措辞与目录不一致，"
                f"或跨页断行导致匹配失败，需人工确认："
                f"{'；'.join(e['title'][:24] for e in unresolved[:5])}",
                [e["title"] for e in unresolved[:5]])
    return f


def check_editorial(d):
    """定稿前校对项：修订痕迹、批注、页眉页脚、域错误、日期格式、标点。"""
    f = []
    add = lambda lv, item, msg, fact=None: f.append(
        {"level": lv, "item": item, "message": msg, "fact": fact})
    ed = d.get("editorial") or {}

    # ① 修订痕迹与批注 —— 带着这些发出去属于事故，级别最高
    tc = ed.get("tracked_changes") or {}
    tot = sum(v for v in tc.values() if isinstance(v, int))
    if tot:
        detail = "、".join(f"{k}{v}处" for k, v in
                          (("插入", tc.get("insertions", 0)), ("删除", tc.get("deletions", 0)),
                           ("移动", tc.get("moves", 0)), ("格式修订", tc.get("format_changes", 0)))
                          if v)
        add("ERROR", "残留修订痕迹",
            f"文档仍带有未接受的修订痕迹（{detail}）。定稿前必须「接受所有修订」并关闭修订模式，"
            f"否则接收方会看到全部修改历史", tc)
    if ed.get("comments"):
        add("ERROR", "残留批注",
            f"文档中还有 {ed['comments']} 条批注未删除，定稿前应清除", ed["comments"])

    # ② 域代码错误（目录/交叉引用失效后会直接印在纸面上）
    fe = ed.get("field_errors") or []
    if fe:
        add("ERROR", "域代码错误残留",
            f"检出 {len(fe)} 处「未找到引用源」类错误文本，说明交叉引用或书签已失效："
            f"{'；'.join(x[:30] for x in fe[:3])}", fe[:3])

    # ③ 页眉页脚一致性（版本号/编号改了但页眉没改，是高发问题）
    for key, label in (("headers", "页眉"), ("footers", "页脚")):
        vals = [v for v in (ed.get(key) or []) if v]
        if len(set(vals)) > 1:
            add("WARN", f"{label}不一致",
                f"各节{label}内容不同（{len(set(vals))} 种）：{'；'.join(sorted(set(vals))[:3])}；"
                f"请确认是否有意为之（改版本号时常漏改其中一处）", sorted(set(vals))[:3])

    # ④ 日期格式混用
    ds = ed.get("date_styles") or {}
    if len(ds) > 1:
        add("WARN", "日期格式不统一",
            f"全文混用 {len(ds)} 种日期写法："
            f"{'、'.join(f'{k}({v}处)' for k, v in sorted(ds.items(), key=lambda x: -x[1]))}", ds)

    # ⑤ 中文文档里的半角标点
    hp = ed.get("half_punct") or []
    if hp:
        n = sum(x.get("count", 1) for x in hp)
        samples = "；".join("{}[{}]".format(x["text"][:26], x["marks"]) for x in hp[:3])
        add("WARN", "中文句中使用半角标点",
            f"{len(hp)} 个段落共 {n} 处中文字后紧跟半角标点（应改用全角）：{samples}", hp[:3])

    # ⑥ 括号不配对
    ub = ed.get("unbalanced") or []
    if ub:
        add("WARN", "括号不配对",
            f"检出 {len(ub)} 处括号数量不匹配："
            f"{'；'.join(x['text'][:28] for x in ub[:3])}", ub[:3])

    # ⑦ 空段落 / 空单元格（结构性编辑残留，方法论第四节）
    if ed.get("consecutive_empty_paragraphs", 0) >= 3:
        add("WARN", "连续空段落",
            f"检出 {ed['consecutive_empty_paragraphs']} 处连续空段落，"
            f"通常是删段后留下的空壳，带列表样式时会渲染成空的项目符号",
            ed["consecutive_empty_paragraphs"])
    if ed.get("empty_table_cells", 0):
        n = ed["empty_table_cells"]
        add("WARN" if n >= 5 else "MANUAL", "表格空单元格",
            f"表格中有 {n} 个空单元格；申报资料中空值一般应写「NA」或「不适用」而非留空", n)

    # ⑧ 文档属性（可能带出上一版作者/单位）
    props = ed.get("properties") or {}
    who = [f"{k}={v}" for k, v in (("作者", props.get("author")),
                                   ("最后修改人", props.get("last_modified_by"))) if v]
    if who:
        add("MANUAL", "文档属性",
            f"文档属性中记录了 {'、'.join(who)}；对外递交前请确认是否需要清除"
            f"（Word：文件→信息→检查问题→检查文档）", props)
    return f


def check_structure(d):
    f = []
    add = lambda lv, item, msg, fact=None: f.append(
        {"level": lv, "item": item, "message": msg, "fact": fact})
    st = d.get("structure") or {}

    # ① 表/图/附录编号是否连续、有无重复
    for kind, nums in (st.get("captions") or {}).items():
        plain = [n for n in nums if re.fullmatch(r"[0-9]+", n)]
        if len(plain) >= 2:
            ints = [int(n) for n in plain]
            dup = [n for n, c in Counter(ints).items() if c > 1]
            if dup:
                add("ERROR", f"{kind}编号重复", f"{kind} {sorted(dup)} 出现了多次", sorted(dup))
            uniq = sorted(set(ints))
            missing = [i for i in range(uniq[0], uniq[-1] + 1) if i not in set(uniq)]
            if missing:
                add("ERROR", f"{kind}编号不连续",
                    f"{kind}编号从 {uniq[0]} 到 {uniq[-1]}，中间缺 {missing}", missing)

    # ② 交叉引用指向的对象是否真实存在
    for kind, refs in (st.get("xrefs") or {}).items():
        have = set((st.get("captions") or {}).get(kind, []))
        alias = {"Table": "表", "Figure": "图", "Appendix": "附录"}.get(kind)
        if alias:
            have |= set((st.get("captions") or {}).get(alias, []))
        miss = [r for r in refs if r not in have]
        if miss and have:
            add("ERROR", f"{kind}交叉引用悬空",
                f"正文引用了 {kind} {miss}，但全文没有找到对应的{kind}标题", miss)

    secs = set(st.get("section_numbers") or [])
    if secs:
        miss = [r for r in (st.get("section_refs") or []) if r not in secs]
        if miss:
            add("WARN", "章节交叉引用存疑",
                f"正文提到「见第 {'、'.join(miss[:8])} 节」，未在标题中找到对应编号（可能是编号格式差异，需人工确认）",
                miss[:8])

    # ③ 同一缩写被定义成了不同全称
    for ab, fulls in (st.get("abbr_defs") or {}).items():
        distinct = _dedupe_fullnames(fulls)
        if len(distinct) > 1:
            add("ERROR", "缩写定义冲突",
                f"缩写「{ab}」在文中对应了多个全称：{'；'.join(distinct[:4])}", distinct[:4])

    # ④ 高频使用但从未定义的缩写
    defined = set((st.get("abbr_defs") or {}).keys())
    undef = [(a, c) for a, c in (st.get("abbr_uses") or []) if a not in defined and c >= 3]
    if undef:
        show = "、".join(f"{a}({c}次)" for a, c in undef[:8])
        add("WARN", "缩写未定义",
            f"以下缩写多次出现但全文未见「全称（缩写）」式定义：{show}；请确认首次出现处是否已定义",
            [a for a, _ in undef[:8]])

    # ⑤ 目录与正文一致性
    f += _check_toc(d)

    # ⑤之二 编辑痕迹与排版细节
    f += check_editorial(d)

    # ⑥ 标题层级跳级
    for sk in (d.get("headings", {}).get("level_skips") or [])[:5]:
        add("WARN", "标题层级跳级",
            f"从 H{sk['from']} 直接跳到 H{sk['to']}：「{sk['text']}」", sk)

    # ⑥ 同一样式内部字体/字号不统一
    for sname, sv in (d.get("styles") or {}).items():
        if len(sv["eastasia"]) > 1:
            add("ERROR", "同级样式不统一",
                f"样式「{sname}」内出现多种中文字体 {sv['eastasia']}", sv["eastasia"])
        if len(sv["sizes_pt"]) > 2:
            add("WARN", "同级样式不统一",
                f"样式「{sname}」内出现多种字号 {sv['sizes_pt']}", sv["sizes_pt"])
    return f


# ══════════════════════════════════════════════════════════════
#  CDE 规则集
# ══════════════════════════════════════════════════════════════

def check_cde(d):
    S = CDE_SPEC
    f = []
    add = lambda lv, item, msg, fact=None: f.append(
        {"level": lv, "item": item, "message": msg, "fact": fact})

    if d["type"] != "docx":
        add("WARN", "文件类型",
            "CDE 字体/字号/行距规则针对可编辑源文件（.docx）判定最准确；PDF 只能做有限核查。"
            "建议用定稿前的 .docx 跑一次本规则集。")
        if not d.get("text_searchable"):
            add("ERROR", "可搜索性", "PDF 文本无法提取（疑似纯扫描件未做 OCR），CDE 要求 PDF 可文本搜索")
        if d["file_size_mb"] > 200:
            add("ERROR", "文件大小",
                f"{d['file_size_mb']}MB，超过 eCTD 技术规范 V1.1 单文件 200MB 上限")
        return f + check_structure(d)

    # ── 中文字体 ──
    dom_cn = d["fonts"]["dominant_eastasia"]
    if dom_cn and not _font_matches(dom_cn, S["cn_font"]):
        add("ERROR", "中文字体", f"正文中文主字体为「{dom_cn}」，规范要求「{S['cn_font']}」", dom_cn)
    others_cn = [n for n, _ in d["fonts"]["eastasia_by_chars"] if not _font_matches(n, S["cn_font"])]
    if others_cn and dom_cn and _font_matches(dom_cn, S["cn_font"]):
        add("WARN", "中文字体", f"另混用了非宋体的中文字体：{'、'.join(others_cn[:6])}", others_cn[:6])

    # ── 西文字体 ──
    dom_en = d["fonts"]["dominant_latin"]
    if dom_en and not _font_matches(dom_en, S["latin_font"]):
        add("ERROR", "西文字体", f"西文主字体为「{dom_en}」，规范要求「{S['latin_font']}」", dom_en)
    others_en = [n for n, _ in d["fonts"]["latin_by_chars"] if not _font_matches(n, S["latin_font"])]
    if others_en and dom_en and _font_matches(dom_en, S["latin_font"]):
        add("WARN", "西文字体", f"另混用了其他西文字体：{'、'.join(others_en[:6])}", others_en[:6])

    # ── 正文字号 ──
    ds = d["fonts"]["dominant_size_pt"]
    if ds is not None:
        if ds < S["body_size_pt"] - 0.01:
            add("ERROR", "正文字号",
                f"正文主字号 {_pt_label(ds)}，低于规范下限 {_pt_label(S['body_size_pt'])}", ds)
        elif ds > S["body_size_pt"] + 0.01:
            add("WARN", "正文字号",
                f"正文主字号 {_pt_label(ds)}，大于约定的 {_pt_label(S['body_size_pt'])}（合规，但请确认是否有意为之）", ds)
    small = [(s, c) for s, c in d["fonts"]["size_pt_by_chars"] if s < S["table_min_size_pt"] - 0.01]
    if small:
        add("ERROR", "字号过小",
            f"文中出现小于表格下限 {_pt_label(S['table_min_size_pt'])} 的字号："
            f"{'、'.join(_pt_label(s) for s, _ in small[:5])}", [s for s, _ in small[:5]])

    # ── 表格字号 ──
    tmin = d["tables"]["min_size_pt"]
    if tmin is not None and tmin < S["table_min_size_pt"] - 0.01:
        add("ERROR", "表格字号",
            f"表格最小字号 {_pt_label(tmin)}，低于规范下限 {_pt_label(S['table_min_size_pt'])}", tmin)

    # ── 行距 ──
    def _ls_desc(x):
        unit = "pt固定值" if x["kind"] == "固定值" else "倍"
        return f"{x['value']}{unit}（{x['paragraphs']}段）"

    ls = d.get("line_spacing") or []
    is_target = lambda x: x["kind"] == "倍数" and abs(x["value"] - S["line_spacing"]) < 0.01
    ok = [x for x in ls if is_target(x)]
    bad = [x for x in ls if not is_target(x)]

    if not ls:
        add("WARN", "行距",
            f"未检出显式行距设置（可能全部继承自模板默认值），请人工确认是否为 {S['line_spacing']} 倍")
    elif not ok:
        add("ERROR", "行距",
            f"未检出 {S['line_spacing']} 倍行距，实际为：{'、'.join(_ls_desc(x) for x in ls[:4])}",
            ls[:4])
    elif bad:
        add("WARN", "行距不统一",
            f"主体为 {S['line_spacing']} 倍行距，但另有 "
            f"{'、'.join(_ls_desc(x) for x in bad[:3])} 等其他设置", bad[:3])

    # ── 字体颜色 ──
    non_black = [(c, n) for c, n in d["fonts"]["color_by_chars"]
                 if c not in ("(继承/自动)", S["font_color_hex"])]
    if non_black:
        total = max(d.get("body_chars") or 1, 1)
        chars = sum(n for _, n in non_black)
        add("ERROR" if chars / total > 0.02 else "WARN", "字体颜色",
            f"检出非黑色文字 {chars} 字：{'、'.join('#' + c for c, _ in non_black[:5])}；规范要求黑色",
            [c for c, _ in non_black[:5]])

    # ── 纸张 ──
    P = S["paper"]
    for i, s in enumerate(d["sections"], 1):
        w, h = s["page_width_cm"], s["page_height_cm"]
        if w and h:
            short, long_ = min(w, h), max(w, h)
            if abs(short - P["width_cm"]) > P["tol_cm"] or abs(long_ - P["height_cm"]) > P["tol_cm"]:
                add("ERROR", "纸张规格",
                    f"第{i}节页面 {w}×{h}cm，规范要求 {P['name']}（21.0×29.7cm）", [w, h])

    # ── 页边距（按纵/横向分别判）──
    for i, s in enumerate(d["sections"], 1):
        rule = S["margins_cm"][s["orientation"]]
        ori = "纵向" if s["orientation"] == "portrait" else "横向"
        for key, label in (("left_cm", "左"), ("right_cm", "右"),
                           ("top_cm", "上"), ("bottom_cm", "下")):
            v, need = s.get(key), rule[key.replace("_cm", "")]
            if v is None:
                continue
            if v < need - 0.05:
                add("ERROR", "页边距",
                    f"第{i}节（{ori}）{label}边距 {v}cm，低于规定的 {need}cm", v)
            elif v > need + 0.3:
                add("WARN", "页边距",
                    f"第{i}节（{ori}）{label}边距 {v}cm，明显大于规定的 {need}cm（合规，但可能是格式不统一）", v)

    ms = {(s["left_cm"], s["right_cm"], s["top_cm"], s["bottom_cm"], s["orientation"])
          for s in d["sections"]}
    by_ori = defaultdict(set)
    for m in ms:
        by_ori[m[4]].add(m[:4])
    for ori, group in by_ori.items():
        if len(group) > 1:
            add("WARN", "页边距不统一",
                f"同为{'纵向' if ori == 'portrait' else '横向'}页面，却存在 {len(group)} 组不同页边距", list(group))

    if d["file_size_mb"] > 200:
        add("ERROR", "文件大小", f"{d['file_size_mb']}MB，超过 eCTD 单文件 200MB 上限", d["file_size_mb"])

    return f + check_structure(d)


# ══════════════════════════════════════════════════════════════
#  FDA 规则集
# ══════════════════════════════════════════════════════════════

def check_fda(d):
    S = FDA_SPEC
    f = []
    add = lambda lv, item, msg, fact=None: f.append(
        {"level": lv, "item": item, "message": msg, "fact": fact})
    ok_fonts = S["latin_fonts_ok"]

    if d["type"] == "docx":
        dom = d["fonts"]["dominant_latin"]
        if dom and not any(_font_matches(dom, e) for e in ok_fonts):
            add("ERROR", "字体", f"正文主字体为「{dom}」，FDA 要求 Times New Roman 或 Arial", dom)
        others = [n for n, _ in d["fonts"]["latin_by_chars"]
                  if not any(_font_matches(n, e) for e in ok_fonts)]
        if others:
            add("WARN", "字体", f"混用了 {len(others)} 种非 Times/Arial 字体：{'、'.join(others[:6])}", others[:6])

        ds = d["fonts"]["dominant_size_pt"]
        if ds is not None and abs(ds - S["body_size_pt"]) > 0.01:
            add("ERROR" if ds < S["table_hard_min_pt"] else "WARN", "字号",
                f"正文主字号 {ds}pt，FDA 建议叙述性正文用 {S['body_size_pt']}pt", ds)

        tmin = d["tables"]["min_size_pt"]
        if tmin is not None:
            if tmin < S["table_hard_min_pt"]:
                add("ERROR", "表格字号",
                    f"表格出现 {tmin}pt，低于可接受下限 {S['table_hard_min_pt']}pt", tmin)
            elif tmin < S["table_min_size_pt"]:
                add("WARN", "表格字号", f"表格最小 {tmin}pt，FDA 建议 ≥{S['table_min_size_pt']}pt", tmin)

        for i, s in enumerate(d["sections"], 1):
            for key, label in (("left_cm", "左"), ("right_cm", "右"),
                               ("top_cm", "上"), ("bottom_cm", "下")):
                v = s.get(key)
                if v is None:
                    continue
                inch = round(v / 2.54, 2)
                hard = S["margin_in"]["hard_left"] if key == "left_cm" else S["margin_in"]["hard_other"]
                if inch < hard:
                    add("ERROR", "页边距", f"第{i}节{label}边距 {inch}in，低于硬性下限 {hard}in", inch)
                elif inch < S["margin_in"]["all"]:
                    add("WARN", "页边距", f"第{i}节{label}边距 {inch}in，FDA 建议四边 ≥1in", inch)
    else:
        ver = d.get("pdf_version")
        if ver:
            try:
                lo, hi = S["pdf_version_range"]
                if not (lo <= float(ver) <= hi):
                    add("ERROR", "PDF 版本", f"版本 {ver}，FDA 接受 {lo}–{hi}", ver)
            except ValueError:
                pass
        if d.get("encrypted"):
            add("ERROR", "安全设置", "PDF 有加密/口令保护，FDA 不接受")
        if d.get("has_javascript"):
            add("ERROR", "动态内容", "PDF 含 JavaScript/OpenAction，FDA 禁止")
        if not d.get("text_searchable"):
            add("ERROR", "可搜索性", "文本无法提取（疑似扫描件未 OCR）")
        if d.get("page_count", 0) >= S["bookmark_required_pages"] and d["bookmarks"]["count"] == 0:
            add("ERROR", "书签", f"{d['page_count']} 页文档无书签，FDA 要求 ≥5 页须有书签")
        if d["bookmarks"]["max_depth"] > S["bookmark_max_depth"]:
            add("WARN", "书签", f"层级 {d['bookmarks']['max_depth']} 层，建议 ≤{S['bookmark_max_depth']} 层")
        emb = d.get("fonts", {}).get("embedding", {})
        if emb.get("not_embedded"):
            add("ERROR", "字体嵌入", f"{emb['not_embedded']} 处字体未嵌入，FDA 要求非标准字体必须嵌入")
        bad = [n for n, _ in d.get("fonts", {}).get("used", [])
               if not any(_font_matches(n, e) for e in ok_fonts)]
        if bad:
            add("WARN", "字体", f"使用了非 Times/Arial 字体：{'、'.join(bad[:6])}", bad[:6])
        m = d.get("min_text_margins_in") or {}
        for key, label, hard in (("left_in", "左", S["margin_in"]["hard_left"]),
                                 ("right_in", "右", S["margin_in"]["hard_other"]),
                                 ("top_in", "上", S["margin_in"]["hard_other"]),
                                 ("bottom_in", "下", S["margin_in"]["hard_other"])):
            if key in m:
                if m[key] < hard:
                    add("ERROR", "页边距", f"{label}侧正文距页边 {m[key]}in，低于硬性下限 {hard}in", m[key])
                elif m[key] < S["margin_in"]["all"]:
                    add("WARN", "页边距", f"{label}侧正文距页边 {m[key]}in，建议 ≥1in", m[key])

    return f + check_structure(d)


PROFILES = {"cde": check_cde, "fda": check_fda}


# ══════════════════════════════════════════════════════════════
#  跨文件一致性
# ══════════════════════════════════════════════════════════════

def compare(datas):
    out = []
    docs = [d for d in datas if d["type"] == "docx"]
    if len(docs) < 2:
        return out
    for key, label in (("dominant_eastasia", "中文主字体"),
                       ("dominant_latin", "西文主字体"),
                       ("dominant_size_pt", "正文主字号")):
        vals = defaultdict(list)
        for d in docs:
            vals[d["fonts"][key]].append(d["file"])
        if len(vals) > 1:
            out.append({"level": "WARN", "item": f"跨文件{label}不一致",
                        "message": "；".join(f"{k}：{'、'.join(v)}" for k, v in vals.items()),
                        "fact": {str(k): v for k, v in vals.items()}})
    mg = defaultdict(list)
    for d in docs:
        if d["sections"]:
            s = d["sections"][0]
            mg[(s["left_cm"], s["right_cm"], s["top_cm"], s["bottom_cm"])].append(d["file"])
    if len(mg) > 1:
        out.append({"level": "WARN", "item": "跨文件页边距不一致",
                    "message": "；".join(f"左右上下{list(k)}cm：{'、'.join(v)}" for k, v in mg.items()),
                    "fact": {str(k): v for k, v in mg.items()}})

    # 同一缩写在不同文件里被定义成不同全称 —— 多文档最常见的一致性问题
    amap = defaultdict(lambda: defaultdict(list))
    for d in docs:
        for ab, fulls in (d.get("structure", {}).get("abbr_defs") or {}).items():
            for full in fulls:
                amap[ab][full].append(d["file"])
    for ab, fulls in amap.items():
        keep = set(_dedupe_fullnames(list(fulls.keys())))
        if len(keep) > 1:
            out.append({"level": "ERROR", "item": "跨文件缩写定义冲突",
                        "message": f"「{ab}」在不同文件里定义不同：" +
                                   "；".join(f"{k}（{'、'.join(v)}）"
                                             for k, v in fulls.items() if k in keep),
                        "fact": {k: v for k, v in fulls.items() if k in keep}})
    return out


# ══════════════════════════════════════════════════════════════
#  输出
# ══════════════════════════════════════════════════════════════

LEVEL_ORDER = {"ERROR": 0, "WARN": 1, "MANUAL": 2}
LEVEL_TAG = {"ERROR": "✗ 错误", "WARN": "⚠ 警告", "MANUAL": "◻ 待人工"}


def print_findings(title, findings):
    print(f"\n{'═' * 70}\n{title}\n{'═' * 70}")
    if not findings:
        print("  ✓ 未发现格式层/结构层问题。")
        return
    for fd in sorted(findings, key=lambda x: LEVEL_ORDER.get(x["level"], 9)):
        print(f"  [{LEVEL_TAG[fd['level']]}] {fd['item']}：{fd['message']}")


def main():
    ap = argparse.ArgumentParser(
        description="临床试验申报资料格式层/结构层核查工具（CDE / FDA）")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p1 = sub.add_parser("extract", help="只提取格式事实，不判定")
    p1.add_argument("files", nargs="+")
    p1.add_argument("--json", action="store_true")

    p2 = sub.add_parser("check", help="按监管机构规则集判定")
    p2.add_argument("files", nargs="+")
    p2.add_argument("--profile", required=True, choices=sorted(PROFILES))
    p2.add_argument("--json", action="store_true")

    p3 = sub.add_parser("compare", help="多份文件之间的格式/缩写一致性")
    p3.add_argument("files", nargs="+")
    p3.add_argument("--json", action="store_true")

    a = ap.parse_args()
    for fp in a.files:
        if not os.path.exists(fp):
            _fail(f"文件不存在：{fp}")
    datas = [extract(fp) for fp in a.files]

    if a.cmd == "extract":
        print(json.dumps(datas if len(datas) > 1 else datas[0],
                         ensure_ascii=False, indent=2, default=str))
        return

    if a.cmd == "compare":
        res = compare(datas)
        if a.json:
            print(json.dumps(res, ensure_ascii=False, indent=2, default=str))
        else:
            print_findings(f"跨文件一致性（{len(datas)} 份）", res)
        return

    checker = PROFILES[a.profile]
    allres = []
    for d in datas:
        allres.append({"file": d["file"], "type": d["type"], "language": d.get("language"),
                       "profile": a.profile, "findings": checker(d)})
    if len(datas) > 1:
        allres.append({"file": "（跨文件）", "type": "-", "profile": a.profile,
                       "findings": compare(datas)})

    if a.json:
        print(json.dumps(allres, ensure_ascii=False, indent=2, default=str))
        return

    for r in allres:
        print_findings(f"{r['file']}　[{a.profile.upper()} 规则集]", r["findings"])
    ne = sum(1 for r in allres for x in r["findings"] if x["level"] == "ERROR")
    nw = sum(1 for r in allres for x in r["findings"] if x["level"] == "WARN")
    print(f"\n{'─' * 70}\n合计：{ne} 项错误、{nw} 项警告。")
    print("提示：本工具只覆盖「格式层 + 结构层」。内容前后矛盾、逻辑自洽性、语言自然性")
    print("      属「文本层」，须按 SKILL.md 的文本层流程另行核查，工具判不了。")


if __name__ == "__main__":
    main()
